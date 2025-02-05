import os
from PyPDF2 import PdfReader
import docx

def parse_resume(file):
    """
    Parse the uploaded resume file (supports .txt, .docx, and .pdf).
    Args:
        file (FileStorage): The uploaded file object.
    Returns:
        str: Extracted text content from the resume.
    """
    try:
        # Handle .txt files
        if file.filename.endswith('.txt'):
            return file.stream.read().decode("utf-8")

        # Handle .docx files
        elif file.filename.endswith('.docx'):
            from docx import Document
            doc = Document(file.stream)
            return "\n".join([p.text for p in doc.paragraphs])

        # Handle .pdf files
        elif file.filename.endswith('.pdf'):
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(file.stream)
            return "\n".join([page.extract_text() for page in pdf_reader.pages])

        else:
            return "Unsupported file format. Please upload a .txt, .docx, or .pdf file."
    except Exception as e:
        return f"Error reading file: {e}"



def parse_txt(file_path):
    """
    Parse a .txt file.
    Args:
        file_path (str): Path to the text file.
    Returns:
        str: Content of the text file.
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def parse_docx(file_path):
    """
    Parse a .docx file.
    Args:
        file_path (str): Path to the Word document.
    Returns:
        str: Content of the Word document.
    """
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


def parse_pdf(file_path):
    """
    Parse a .pdf file.
    Args:
        file_path (str): Path to the PDF file.
    Returns:
        str: Content of the PDF file.
    """
    reader = PdfReader(file_path)
    return "\n".join([page.extract_text() for page in reader.pages])
